import os

from docx import Document
from docx.shared import Cm, Inches

from paper import Paper
from word import WpsWord

test_docx = 'test2.docx'
document = Document()

# 文档页边距设置
sec = document.sections
# 获取、设置页面边距
sec0 = sec[0]  # 获取章节对象
# 获取页面边距值：（单位为像素）
print('左边距：', sec0.left_margin)
# 左边距： 1143000
print('右边距：', sec0.right_margin)
# 右边距： 1143000
print('上边距：', sec0.top_margin)


# 上边距： 914400
# print('下边距：', sec0.bottom_margin)
# 下边距： 914400

def tempA_B():
    w = 15.3
    h = 11
    png = 'image-filename.png'
    p1 = Paper(png, h, w)
    document.add_picture(p1.path, width=Cm(p1.w), height=Cm(p1.h))
    p1 = Paper(png, h, w)
    document.add_picture(p1.path, width=Cm(p1.w), height=Cm(p1.h))


def tempA_BC():
    w = 15.3
    h = 11
    png = 'image-filename.png'
    p1 = Paper(png, h, w)
    document.add_picture(p1.path, width=Cm(p1.w), height=Cm(p1.h))

    tables = document.add_table(rows=1, cols=2)
    p2 = Paper(png, h, w / 2)
    # for i in range(2):
    #     run = tables.cell(0, i).paragraphs[0].add_run()
    #     run.add_picture(p2.path, width=Cm(p2.w), height=Cm(p2.h))
    #
    run = tables.cell(0, 0).paragraphs[0].add_run()
    run.add_picture(p2.path, width=Cm(p2.w), height=Cm(p2.h))

    paragraphs_ = tables.cell(0, 1).paragraphs[0]
    run = paragraphs_.add_run()
    paragraphs_.top_margin= Inches(0.5)
    run.add_picture(p2.path, width=Cm(p2.w-1), height=Cm(p2.h-1))


def margin1(sec0):
    sec0.left_margin = Inches(1.2)
    sec0.right_margin = Inches(1.2)
    sec0.top_margin = Inches(1)
    # sec0.bottom_margin = Inches(1)
    print('左边距：', sec0.left_margin)
    # 左边距： 1143000
    print('右边距：', sec0.right_margin)
    # 右边距： 1143000
    print('上边距：', sec0.top_margin)
    # 上边距： 914400
    # print('下边距：', sec0.bottom_margin)


# margin1(sec0)
tempA_BC()

if os.path.exists(test_docx):
    os.remove(test_docx)
document.save(test_docx)

root = "C:\\Users\\Administrator\\IdeaProjects\\stop-motion-printing\\"
WpsWord().open(root + test_docx)
# test_docx = 'C:\\Users\\Administrator\\IdeaProjects\\stop-motion-printing\\test1.docx'
