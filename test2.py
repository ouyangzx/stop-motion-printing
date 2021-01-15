import math

from docx import Document
from docx.shared import Inches

doc = Document()
tables = doc.add_table(rows=1, cols=3)
num = math.ceil(3 / 3)
png = 'image-filename.png'

for i in range(3):
    run = tables.cell(0, i).paragraphs[0].add_run()
    run.add_picture(png, width=Inches(2), height=Inches(2.5))
doc.save('1.docx')
